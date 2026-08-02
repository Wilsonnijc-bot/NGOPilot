from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .errors import VisitNoteAgentError


def render_docx_from_template(
    template_path: str,
    template_contract: dict[str, Any],
    generated_content: dict[str, Any],
    output_path: str,
) -> None:
    try:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(template_path, output_path)
        document = Document(output_path)
        paragraph_map = _body_paragraph_map(document)
        table_cell_map = _table_cell_map(document)
        header_footer_map = _header_footer_paragraph_map(document)

        for slot in template_contract.get("dynamic_slots", []):
            slot_id = slot.get("slot_id")
            source_block_id = slot.get("source_block_id")
            if not slot_id or not source_block_id or slot_id not in generated_content:
                continue
            value = generated_content[slot_id]
            if source_block_id.startswith("tbl_"):
                cell = table_cell_map.get(source_block_id)
                if cell is not None:
                    _replace_paragraph_text(cell.paragraphs[0], _value_to_text(value))
            else:
                paragraph = paragraph_map.get(source_block_id) or header_footer_map.get(
                    source_block_id
                )
                if paragraph is not None:
                    if _should_replace_as_list_block(slot, value, paragraph):
                        _replace_list_block(slot, value, paragraph_map, header_footer_map)
                    else:
                        _replace_paragraph_text(paragraph, _value_to_text(value))

        document.save(output_path)
    except Exception as exc:  # pragma: no cover - exact python-docx errors vary
        raise VisitNoteAgentError("DOCX template rendering failed") from exc


def _body_paragraph_map(document: DocumentObject) -> dict[str, Paragraph]:
    mapping: dict[str, Paragraph] = {}
    paragraph_index = 0
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        paragraph_index += 1
        mapping[f"p_{paragraph_index:03d}"] = paragraph
    return mapping


def _table_cell_map(document: DocumentObject) -> dict[str, Any]:
    mapping: dict[str, Any] = {}
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for column_index, cell in enumerate(row.cells, start=1):
                mapping[f"tbl_{table_index:03d}_r{row_index:03d}_c{column_index:03d}"] = cell
    return mapping


def _header_footer_paragraph_map(document: DocumentObject) -> dict[str, Paragraph]:
    mapping: dict[str, Paragraph] = {}
    for section_index, section in enumerate(document.sections, start=1):
        for container_name, part in (("header", section.header), ("footer", section.footer)):
            paragraph_index = 0
            for paragraph in part.paragraphs:
                if not paragraph.text.strip():
                    continue
                paragraph_index += 1
                mapping[f"{container_name}_{section_index:03d}_p{paragraph_index:03d}"] = paragraph
    return mapping


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def _replace_list_block(
    slot: dict[str, Any],
    value: Any,
    paragraph_map: dict[str, Paragraph],
    header_footer_map: dict[str, Paragraph],
) -> None:
    lookup = {**paragraph_map, **header_footer_map}
    source_block_id = slot.get("source_block_id")
    if not source_block_id:
        return
    block_ids = _list_block_ids(slot, source_block_id, paragraph_map, header_footer_map)
    paragraphs = [lookup[block_id] for block_id in block_ids if block_id in lookup]
    if not paragraphs:
        return

    items = _value_to_list_items(value)
    for paragraph, item in zip(paragraphs, items):
        _replace_paragraph_text(paragraph, item)

    if len(items) < len(paragraphs):
        for paragraph in paragraphs[len(items) :]:
            _delete_paragraph(paragraph)
        return

    insertion_point = paragraphs[-1]
    style_source = paragraphs[0]
    for item in items[len(paragraphs) :]:
        insertion_point = _insert_paragraph_after(insertion_point, item, style_source)


def _list_block_ids(
    slot: dict[str, Any],
    source_block_id: str,
    paragraph_map: dict[str, Paragraph],
    header_footer_map: dict[str, Paragraph],
) -> list[str]:
    source_block_ids = slot.get("source_block_ids")
    if isinstance(source_block_ids, list) and source_block_ids:
        return [str(block_id) for block_id in source_block_ids if block_id]
    if source_block_id in paragraph_map:
        return _infer_contiguous_list_ids(source_block_id, paragraph_map)
    return [source_block_id] if source_block_id in header_footer_map else []


def _infer_contiguous_list_ids(
    source_block_id: str, paragraph_map: dict[str, Paragraph]
) -> list[str]:
    block_ids = list(paragraph_map)
    start = block_ids.index(source_block_id)
    source_paragraph = paragraph_map[source_block_id]
    source_style = source_paragraph.style.name if source_paragraph.style else ""
    inferred = [source_block_id]
    for block_id in block_ids[start + 1 :]:
        paragraph = paragraph_map[block_id]
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name != source_style or not _is_list_paragraph(paragraph):
            break
        inferred.append(block_id)
    return inferred


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if "list" in style_name or "bullet" in style_name or "number" in style_name:
        return True
    ppr = paragraph._p.pPr
    return bool(ppr is not None and ppr.numPr is not None) or _looks_like_list_item_text(
        paragraph.text
    )


def _should_replace_as_list_block(
    slot: dict[str, Any], value: Any, paragraph: Paragraph
) -> bool:
    if slot.get("replacement_unit") == "list_block":
        return True
    return _is_list_paragraph(paragraph) and len(_value_to_list_items(value)) > 1


def _looks_like_list_item_text(text: str) -> bool:
    return bool(re.match(r"^\s*(?:[-*•]|[0-9０-９]+[.)、．])\s*", text))


def _insert_paragraph_after(
    paragraph: Paragraph, text: str, style_source: Paragraph
) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.style = style_source.style
    _replace_paragraph_text(new_paragraph, text)
    return new_paragraph


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _value_to_list_items(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value]
    lines = [line.strip() for line in str(value).splitlines() if line.strip()]
    return lines or [str(value)]


def _value_to_text(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    return str(value)
