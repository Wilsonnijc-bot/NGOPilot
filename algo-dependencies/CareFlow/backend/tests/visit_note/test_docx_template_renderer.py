from docx import Document

from app.services.visit_note_agent.docx_template_renderer import render_docx_from_template


def test_replaces_paragraph_without_touching_fixed_heading(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_heading("探訪紀錄", level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run("舊內容").bold = True
    doc.save(template)

    contract = {
        "dynamic_slots": [
            {
                "slot_id": "summary",
                "source_block_id": "p_002",
                "replacement_unit": "whole_paragraph",
            }
        ]
    }

    render_docx_from_template(str(template), contract, {"summary": "新內容"}, str(output))

    result = Document(output)
    assert result.paragraphs[0].text == "探訪紀錄"
    assert result.paragraphs[1].text == "新內容"
    assert result.paragraphs[1].runs[0].bold is True


def test_renders_generated_cantonese_paragraph_without_mojibake(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_heading("探訪紀錄", level=1)
    doc.add_paragraph("長者姓名：舊姓名")
    doc.save(template)

    contract = {
        "dynamic_slots": [
            {
                "slot_id": "elder_name",
                "source_block_id": "p_002",
                "replacement_unit": "whole_paragraph",
            }
        ]
    }

    render_docx_from_template(
        str(template), contract, {"elder_name": "長者姓名：測試姓名"}, str(output)
    )

    result = Document(output)
    assert result.paragraphs[1].text == "長者姓名：測試姓名"
    assert "é" not in result.paragraphs[1].text


def test_replaces_table_cell_without_changing_table_shape(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "欄位"
    table.cell(0, 1).text = "內容"
    table.cell(1, 0).text = "健康"
    table.cell(1, 1).text = "舊健康內容"
    doc.save(template)

    contract = {
        "dynamic_slots": [
            {
                "slot_id": "health",
                "source_block_id": "tbl_001_r002_c002",
                "replacement_unit": "whole_cell",
            }
        ]
    }

    render_docx_from_template(str(template), contract, {"health": "胃口一般"}, str(output))

    result = Document(output)
    table = result.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.cell(1, 0).text == "健康"
    assert table.cell(1, 1).text == "胃口一般"


def test_renders_generated_cantonese_table_cell_without_mojibake(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "長者姓名"
    table.cell(0, 1).text = "舊姓名"
    doc.save(template)

    contract = {
        "dynamic_slots": [
            {
                "slot_id": "elder_name",
                "source_block_id": "tbl_001_r001_c002",
                "replacement_unit": "whole_cell",
            }
        ]
    }

    render_docx_from_template(str(template), contract, {"elder_name": "測試姓名"}, str(output))

    result = Document(output)
    assert result.tables[0].cell(0, 0).text == "長者姓名"
    assert result.tables[0].cell(0, 1).text == "測試姓名"
    assert "é" not in result.tables[0].cell(0, 1).text


def test_replaces_list_item_and_ignores_fixed_blocks(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("固定說明")
    item = doc.add_paragraph("舊跟進")
    item.style = "List Bullet"
    doc.save(template)

    contract = {
        "fixed_blocks": [{"block_id": "p_001", "replace": False}],
        "dynamic_slots": [
            {
                "slot_id": "follow_up",
                "source_block_id": "p_002",
                "replacement_unit": "list_block",
            }
        ],
    }

    render_docx_from_template(
        str(template), contract, {"follow_up": ["下次再了解覆診安排。"]}, str(output)
    )

    result = Document(output)
    assert result.paragraphs[0].text == "固定說明"
    assert result.paragraphs[1].text == "下次再了解覆診安排。"
    assert "List" in result.paragraphs[1].style.name


def test_replaces_whole_cantonese_list_block_and_clears_stale_items(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_heading("跟進事項", level=1)
    for text in ("轉介家居安全評估服務。", "提醒覆診安排。", "安排義工探訪。"):
        item = doc.add_paragraph(text)
        item.style = "List Bullet"
    doc.save(template)

    contract = {
        "dynamic_slots": [
            {
                "slot_id": "follow_up",
                "source_block_id": "p_002",
                "source_block_ids": ["p_002", "p_003", "p_004"],
                "replacement_unit": "list_block",
            }
        ]
    }

    render_docx_from_template(
        str(template),
        contract,
        {"follow_up": ["下次探訪再了解覆診安排。"]},
        str(output),
    )

    result = Document(output)
    texts = [paragraph.text for paragraph in result.paragraphs]
    assert "下次探訪再了解覆診安排。" in texts
    assert "轉介家居安全評估服務。" not in texts
    assert "提醒覆診安排。" not in texts
    assert "安排義工探訪。" not in texts
    assert "é" not in "\n".join(texts)


def test_multiline_numbered_replacement_replaces_plain_text_numbered_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_heading("跟進計劃", level=1)
    doc.add_paragraph("1. 協助參與社交小組。")
    doc.add_paragraph("2. 轉介家居安全評估服務。")
    doc.add_paragraph("3. 兩星期後電話跟進陳伯。")
    doc.save(template)

    contract = {
        "dynamic_slots": [
            {
                "slot_id": "follow_up",
                "source_block_id": "p_002",
                "replacement_unit": "whole_paragraph",
            }
        ]
    }

    render_docx_from_template(
        str(template),
        contract,
        {
            "follow_up": "\n".join(
                [
                    "1. 協助黃婆婆報名參加長者茶聚。",
                    "2. 安排義工陪同前往中心。",
                ]
            )
        },
        str(output),
    )

    result = Document(output)
    texts = [paragraph.text for paragraph in result.paragraphs if paragraph.text.strip()]
    assert texts == [
        "跟進計劃",
        "1. 協助黃婆婆報名參加長者茶聚。",
        "2. 安排義工陪同前往中心。",
    ]
