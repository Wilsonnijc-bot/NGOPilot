from docx import Document

from app.services.visit_note_agent.docx_structural_extractor import extract_docx_structural_map


def test_extracts_ordered_paragraph_table_header_and_footer_blocks(tmp_path):
    template = tmp_path / "template.docx"
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "機構頁首"
    section.footer.paragraphs[0].text = "機構頁尾"
    doc.add_heading("探訪紀錄", level=1)
    doc.add_paragraph("服務使用者姓名：陳太")
    list_item = doc.add_paragraph("下次再了解覆診安排。")
    list_item.style = "List Bullet"
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "範疇"
    table.cell(0, 1).text = "內容"
    table.cell(1, 0).text = "健康"
    table.cell(1, 1).text = "胃口一般"
    doc.save(template)

    structural_map = extract_docx_structural_map(str(template))
    blocks = structural_map["blocks"]
    by_id = {block["block_id"]: block for block in blocks}

    assert [block["block_id"] for block in blocks[:3]] == ["p_001", "p_002", "p_003"]
    assert by_id["p_001"]["text"] == "探訪紀錄"
    assert by_id["p_002"]["nearest_heading"] == "探訪紀錄"
    assert by_id["p_002"]["location"]["paragraph_index"] == 2
    assert by_id["p_003"]["is_list_item"] is True
    assert by_id["p_003"]["type"] == "list_item"
    assert by_id["tbl_001_r001_c002"]["left_neighbor"] == "範疇"
    assert by_id["tbl_001_r002_c002"]["top_neighbor"] == "內容"
    assert by_id["tbl_001_r002_c002"]["nearest_heading"] == "探訪紀錄"
    assert by_id["header_001_p001"]["container"] == "header"
    assert by_id["footer_001_p001"]["container"] == "footer"


def test_structural_map_keeps_template_metadata(tmp_path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("內容")
    doc.save(template)

    structural_map = extract_docx_structural_map(str(template))

    assert structural_map["template_path"] == str(template)
    assert structural_map["format"] == "docx"
    assert structural_map["block_count"] == 1
