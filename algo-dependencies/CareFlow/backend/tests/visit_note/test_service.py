import os
import re
from threading import Event

import pytest
from docx import Document

import app.services.visit_note_agent.service as service
from app.services.visit_note_agent.errors import VisitNoteAgentError
from app.services.visit_note_agent.service import generate_visit_case_note


def test_generate_visit_case_note_rejects_missing_audio(tmp_path):
    template = tmp_path / "template.docx"
    template.write_text("placeholder", encoding="utf-8")

    with pytest.raises(VisitNoteAgentError) as exc:
        generate_visit_case_note(str(tmp_path / "missing.mp3"), str(template), str(tmp_path))

    assert "Audio file not found" in str(exc.value)


def test_generate_visit_case_note_orchestrates_without_logging_transcript(
    tmp_path, monkeypatch, capsys
):
    audio = tmp_path / "visit.mp3"
    template = tmp_path / "template.docx"
    working = tmp_path / "work" / "normalized_template.docx"
    output_dir = tmp_path / "outputs"
    audio.write_bytes(b"fake audio")
    template.write_bytes(b"fake template")
    calls = []

    monkeypatch.setattr(
        "app.services.visit_note_agent.service.transcribe_audio",
        lambda path: calls.append(("transcribe", path)) or "敏感逐字稿內容",
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service._prepare_template_contract",
        lambda path, temp_dir, model=None: calls.append(("template_branch", path, model))
        or (
            str(working),
            {"blocks": [{"block_id": "p_001"}]},
            {"dynamic_slots": [{"slot_id": "summary", "source_block_id": "p_001"}]},
        ),
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.generate_slot_content",
        lambda transcript, contract, model=None, mode=None: calls.append(("slot_content", transcript))
        or {"summary": "新內容"},
    )

    def fake_render(template_path, contract, generated_content, output_path):
        calls.append(("render", template_path, generated_content))
        with open(output_path, "wb") as handle:
            handle.write(b"docx")

    monkeypatch.setattr("app.services.visit_note_agent.service.render_docx_from_template", fake_render)

    output_path = generate_visit_case_note(
        str(audio), str(template), str(output_dir), model="custom-model"
    )

    captured = capsys.readouterr()
    assert "敏感逐字稿內容" not in captured.out
    assert ("transcribe", str(audio)) in calls
    assert ("template_branch", str(template), "custom-model") in calls
    assert calls.index(("slot_content", "敏感逐字稿內容")) > calls.index(
        ("transcribe", str(audio))
    )
    assert calls.index(("slot_content", "敏感逐字稿內容")) > calls.index(
        ("template_branch", str(template), "custom-model")
    )
    assert calls[-1] == ("render", str(working), {"summary": "新內容"})
    assert output_path.startswith(str(output_dir))
    assert os.path.basename(output_path).startswith("visit_note_")
    assert output_path.endswith(".docx")
    assert re.search(r"visit_note_\d{8}_\d{6}\.docx$", os.path.basename(output_path))
    assert os.path.exists(output_path)


def test_docx_template_uses_clone_and_fill_path(tmp_path, monkeypatch, capsys):
    audio = tmp_path / "visit.mp3"
    template = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"
    audio.write_bytes(b"fake audio")
    template.write_bytes(b"fake template")
    calls = []

    monkeypatch.setattr(
        "app.services.visit_note_agent.service.normalize_template_to_docx",
        lambda path, working_dir: calls.append(("normalize", path)) or str(template),
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.transcribe_audio",
        lambda path: calls.append("transcribe") or "敏感逐字稿內容",
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.extract_docx_structural_map",
        lambda path: calls.append(("structural", path)) or {"blocks": [{"block_id": "p_001"}]},
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.analyze_template_contract",
        lambda structural_map, model=None: calls.append(("analyze", model))
        or {"dynamic_slots": [{"slot_id": "summary", "source_block_id": "p_001"}]},
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.generate_slot_content",
        lambda transcript, contract, model=None, mode=None: calls.append(("content", transcript, model))
        or {"summary": "新內容"},
    )

    def fake_render(template_path, contract, generated_content, output_path):
        calls.append(("render", template_path, generated_content))
        with open(output_path, "wb") as handle:
            handle.write(b"docx")

    monkeypatch.setattr("app.services.visit_note_agent.service.render_docx_from_template", fake_render)

    output_path = generate_visit_case_note(
        str(audio), str(template), str(output_dir), model="slot-model"
    )

    captured = capsys.readouterr()
    assert "敏感逐字稿內容" not in captured.out
    assert ("normalize", str(template)) in calls
    assert "transcribe" in calls
    assert ("structural", str(template)) in calls
    assert ("analyze", "slot-model") in calls
    assert calls.index(("content", "敏感逐字稿內容", "slot-model")) > calls.index(
        ("analyze", "slot-model")
    )
    assert calls[-1] == ("render", str(template), {"summary": "新內容"})
    assert output_path.endswith(".docx")
    assert os.path.exists(output_path)


def test_doc_template_converts_then_uses_clone_and_fill_path(tmp_path, monkeypatch):
    audio = tmp_path / "visit.mp3"
    template = tmp_path / "template.doc"
    normalized = tmp_path / "work" / "normalized_template.docx"
    output_dir = tmp_path / "outputs"
    audio.write_bytes(b"fake audio")
    template.write_bytes(b"doc bytes")
    calls = []

    monkeypatch.setattr(
        "app.services.visit_note_agent.service.normalize_template_to_docx",
        lambda path, working_dir: calls.append(("normalize", path)) or str(normalized),
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.transcribe_audio",
        lambda path: calls.append(("transcribe", path)) or "逐字稿",
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.extract_docx_structural_map",
        lambda path: calls.append(("structural", path)) or {"blocks": [{"block_id": "p_001"}]},
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.analyze_template_contract",
        lambda structural_map, model=None: calls.append(("analyze", model))
        or {"dynamic_slots": [{"slot_id": "summary", "source_block_id": "p_001"}]},
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.generate_slot_content",
        lambda transcript, contract, model=None, mode=None: calls.append(("content", transcript))
        or {"summary": "新內容"},
    )

    def fake_render(template_path, contract, generated_content, output_path):
        calls.append(("render", template_path, generated_content))
        with open(output_path, "wb") as handle:
            handle.write(b"docx")

    monkeypatch.setattr("app.services.visit_note_agent.service.render_docx_from_template", fake_render)

    output_path = generate_visit_case_note(str(audio), str(template), str(output_dir))

    assert ("normalize", str(template)) in calls
    assert ("transcribe", str(audio)) in calls
    assert ("structural", str(normalized)) in calls
    assert ("analyze", None) in calls
    assert calls.index(("content", "逐字稿")) > calls.index(("transcribe", str(audio)))
    assert calls.index(("content", "逐字稿")) > calls.index(("analyze", None))
    assert calls[-1] == ("render", str(normalized), {"summary": "新內容"})
    assert output_path.endswith(".docx")


def test_pdf_template_fails_without_old_fallback_or_renderer(tmp_path, monkeypatch):
    audio = tmp_path / "visit.mp3"
    template = tmp_path / "template.pdf"
    output_dir = tmp_path / "outputs"
    audio.write_bytes(b"fake audio")
    template.write_bytes(b"%PDF fake")

    monkeypatch.setattr(
        "app.services.visit_note_agent.service.transcribe_audio",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("STT not used for PDF")),
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.render_docx_from_template",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("renderer not used for PDF")),
    )

    with pytest.raises(VisitNoteAgentError) as exc:
        generate_visit_case_note(str(audio), str(template), str(output_dir))

    assert str(exc.value) == (
        "PDF templates are no longer supported. Please provide a .docx or .doc template."
    )


def test_transcription_and_template_contract_branches_run_in_parallel(
    tmp_path, monkeypatch
):
    audio = tmp_path / "visit.mp3"
    template = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"
    audio.write_bytes(b"fake audio")
    template.write_bytes(b"fake template")
    transcribe_started = Event()
    template_started = Event()
    transcript_ready = Event()
    template_ready = Event()
    calls = []

    def fake_transcribe(path):
        calls.append(("transcribe_start", path))
        transcribe_started.set()
        assert template_started.wait(1)
        transcript_ready.set()
        calls.append("transcribe_end")
        return "逐字稿"

    def fake_prepare(path, temp_dir, model=None):
        calls.append(("template_start", path, model))
        template_started.set()
        assert transcribe_started.wait(1)
        template_ready.set()
        calls.append("template_end")
        return (
            str(template),
            {"blocks": [{"block_id": "p_001"}]},
            {"dynamic_slots": [{"slot_id": "summary", "source_block_id": "p_001"}]},
        )

    def fake_generate_slot_content(transcript, contract, model=None, mode=None):
        assert transcript_ready.is_set()
        assert template_ready.is_set()
        calls.append(("content", transcript, model))
        return {"summary": "新內容"}

    def fake_render(template_path, contract, generated_content, output_path):
        calls.append(("render", template_path, generated_content))
        with open(output_path, "wb") as handle:
            handle.write(b"docx")

    monkeypatch.setattr("app.services.visit_note_agent.service.transcribe_audio", fake_transcribe)
    monkeypatch.setattr(
        "app.services.visit_note_agent.service._prepare_template_contract", fake_prepare
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.generate_slot_content", fake_generate_slot_content
    )
    monkeypatch.setattr("app.services.visit_note_agent.service.render_docx_from_template", fake_render)

    output_path = generate_visit_case_note(
        str(audio), str(template), str(output_dir), model="slot-model"
    )

    assert calls[:2] in (
        [("transcribe_start", str(audio)), ("template_start", str(template), "slot-model")],
        [("template_start", str(template), "slot-model"), ("transcribe_start", str(audio))],
    )
    assert calls.index(("content", "逐字稿", "slot-model")) > calls.index("transcribe_end")
    assert calls.index(("content", "逐字稿", "slot-model")) > calls.index("template_end")
    assert calls[-1] == ("render", str(template), {"summary": "新內容"})
    assert output_path.endswith(".docx")


def test_service_clone_fill_output_keeps_cantonese_and_removes_stale_list_items(
    tmp_path, monkeypatch, capsys
):
    audio = tmp_path / "visit.mp3"
    template = tmp_path / "template.docx"
    output_dir = tmp_path / "outputs"
    audio.write_bytes(b"fake audio")

    doc = Document()
    doc.add_heading("長者個案面談紀錄", level=1)
    doc.add_paragraph("長者姓名：舊姓名")
    doc.add_heading("跟進事項", level=2)
    for text in ("轉介家居安全評估服務。", "提醒覆診安排。"):
        item = doc.add_paragraph(text)
        item.style = "List Bullet"
    doc.save(template)

    contract = {
        "fixed_blocks": [{"block_id": "p_001", "replace": False}],
        "dynamic_slots": [
            {
                "slot_id": "elder_name",
                "source_block_id": "p_002",
                "replacement_unit": "whole_paragraph",
            },
            {
                "slot_id": "follow_up",
                "source_block_id": "p_004",
                "source_block_ids": ["p_004", "p_005"],
                "replacement_unit": "list_block",
            },
        ],
        "rules": ["Only replace dynamic slots."],
    }
    generated = {
        "elder_name": "長者姓名：測試姓名",
        "follow_up": ["下次探訪再了解覆診安排。"],
    }
    render_calls = []
    real_render = service.render_docx_from_template

    monkeypatch.setattr(
        "app.services.visit_note_agent.service.transcribe_audio",
        lambda path: "敏感逐字稿內容",
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.analyze_template_contract",
        lambda structural_map, model=None: contract,
    )
    monkeypatch.setattr(
        "app.services.visit_note_agent.service.generate_slot_content",
        lambda transcript, template_contract, model=None, mode=None: generated,
    )

    def tracking_render(template_path, template_contract, generated_content, output_path):
        render_calls.append((template_path, generated_content))
        real_render(template_path, template_contract, generated_content, output_path)

    monkeypatch.setattr("app.services.visit_note_agent.service.render_docx_from_template", tracking_render)

    output_path = generate_visit_case_note(str(audio), str(template), str(output_dir))

    captured = capsys.readouterr()
    result = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in result.paragraphs)
    assert "敏感逐字稿內容" not in captured.out
    assert "長者姓名：測試姓名" in text
    assert "下次探訪再了解覆診安排。" in text
    assert "轉介家居安全評估服務。" not in text
    assert "提醒覆診安排。" not in text
    assert not any(marker in text for marker in ("é", "æ", "ï¼", "ã"))
    assert render_calls
    assert not hasattr(service, "write_docx")
    assert not hasattr(service, "generate_note")
